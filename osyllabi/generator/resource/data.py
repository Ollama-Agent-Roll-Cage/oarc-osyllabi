"""
Data file extraction for curriculum generation.

This module provides classes for extracting and analyzing content from data files
such as CSV, JSON, and other structured data formats.
"""
import json
import csv
import importlib.util
from pathlib import Path
from typing import List, Dict, Any, Optional, Set, Tuple

from osyllabi.utils.log import log
from osyllabi.generator.resource.extractor import ContentExtractorABC

# Optional imports with fallbacks
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    pd = None
    PANDAS_AVAILABLE = False

try:
    import yaml
    YAML_AVAILABLE = True
except ImportError:
    yaml = None
    YAML_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False


class DataFileExtractor(ContentExtractorABC):
    """
    Extracts and analyzes content from data files.
    
    This class provides functionality to extract meaningful information from
    data files, including format detection, structure analysis, and summarization.
    """
    
    def __init__(self):
        """Initialize the data file extractor."""
        # Format handlers
        self.format_handlers = {
            '.json': self._handle_json,
            '.csv': self._handle_csv,
            '.tsv': self._handle_csv,
            '.yml': self._handle_yaml,
            '.yaml': self._handle_yaml,
            '.docx': self._handle_docx
        }
        
        # Store availability of optional dependencies
        self.pandas_available = PANDAS_AVAILABLE
        self.yaml_available = YAML_AVAILABLE
        self.docx_available = DOCX_AVAILABLE
    
    def extract(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract and analyze content from a data file.
        
        Args:
            file_path: Path to the data file
            
        Returns:
            Dictionary with extracted content and analysis
        """
        extension = file_path.suffix.lower()
        
        # Use specific handler if available
        if extension in self.format_handlers:
            try:
                return self.format_handlers[extension](file_path)
            except Exception as e:
                log.warning(f"Error processing {extension} file {file_path}: {e}")
        
        # Fall back to basic text extraction
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
                
            return {
                "title": file_path.name,
                "content": content,
                "content_type": "text",
                "extension": extension,
                "metadata": {
                    "size_bytes": file_path.stat().st_size
                }
            }
        except Exception as e:
            log.warning(f"Error reading file {file_path}: {e}")
            return {
                "title": file_path.name,
                "content": f"[Error reading file: {e}]",
                "content_type": "error",
                "metadata": {
                    "size_bytes": file_path.stat().st_size,
                    "error": str(e)
                }
            }
    
    def supports(self, resource: Any) -> bool:
        """Check if this extractor supports the resource."""
        if isinstance(resource, Path) and resource.is_file():
            extension = resource.suffix.lower()
            return extension in self.format_handlers
        return False
    
    def _handle_json(self, file_path: Path) -> Dict[str, Any]:
        """Handle JSON file extraction."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                data = json.load(f)
                
            # Analyze JSON structure
            structure = self._analyze_json(data)
            summary = self._create_json_summary(data, structure)
            
            # Format for display
            formatted_json = json.dumps(data, indent=2, ensure_ascii=False)
            
            return {
                "title": file_path.name,
                "content": f"{summary}\n\n```json\n{formatted_json}\n```",
                "content_type": "json",
                "extension": file_path.suffix,
                "metadata": {
                    "size_bytes": file_path.stat().st_size,
                    "structure": structure
                }
            }
        except json.JSONDecodeError as e:
            # Handle invalid JSON
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
                
            return {
                "title": file_path.name,
                "content": f"Invalid JSON content:\n{content}",
                "content_type": "text",
                "extension": file_path.suffix,
                "metadata": {
                    "size_bytes": file_path.stat().st_size,
                    "error": f"JSON parse error: {e}"
                }
            }
    
    def _handle_csv(self, file_path: Path) -> Dict[str, Any]:
        """Handle CSV file extraction."""
        try:
            if self.pandas_available:
                # Use pandas for CSV handling
                df = pd.read_csv(file_path)
                content = f"# CSV Data Analysis\n\n"
                content += f"## Statistics\n"
                content += f"- Rows: {len(df)}\n"
                content += f"- Columns: {len(df.columns)}\n\n"
                content += f"## Preview\n```\n{df.head().to_string()}\n```\n\n"
                
                # Basic statistics
                stats = {}
                numeric_cols = df.select_dtypes(include=['int64', 'float64']).columns
                for col in numeric_cols:
                    stats[col] = {
                        'min': df[col].min(),
                        'max': df[col].max(),
                        'mean': df[col].mean(),
                        'median': df[col].median()
                    }
                
                return {
                    "title": file_path.name,
                    "content": content,
                    "content_type": "csv",
                    "metadata": {
                        "row_count": len(df),
                        "column_count": len(df.columns),
                        "headers": list(df.columns),
                        "statistics": stats
                    }
                }
            else:
                # Fallback to basic CSV handling
                with open(file_path, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    rows = list(reader)
                    headers = rows[0] if rows else []
                    
                    # Create content with statistics
                    content = "# CSV Data Analysis\n\n"
                    content += f"## Statistics\n"
                    content += f"- Rows: {len(rows)}\n"
                    content += f"- Columns: {len(headers)}\n\n"
                    content += "\n".join([",".join(row) for row in rows])
                    
                    return {
                        "title": file_path.name,
                        "content": content,
                        "content_type": "csv",
                        "metadata": {
                            "row_count": len(rows),
                            "column_count": len(headers),
                            "headers": headers
                        }
                    }
        except Exception as e:
            log.warning(f"Error processing CSV {file_path}: {e}")
            return {
                "title": file_path.name,
                "content": f"Error processing CSV file: {e}",
                "content_type": "text",
                "metadata": {"error": str(e)}
            }
    
    def _handle_yaml(self, file_path: Path) -> Dict[str, Any]:
        """Handle YAML file extraction."""
        try:
            # Use PyYAML if available
            if self.yaml_available:
                import yaml
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    data = yaml.safe_load(f)
                    
                # Re-format as JSON for display (more universal)
                return {
                    "title": file_path.name,
                    "content": f"YAML content (converted to JSON):\n```json\n{json.dumps(data, indent=2)}\n```",
                    "content_type": "yaml",
                    "extension": file_path.suffix,
                    "metadata": {
                        "size_bytes": file_path.stat().st_size
                    }
                }
            else:
                # PyYAML not available, just show raw content
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()
                    
                return {
                    "title": file_path.name,
                    "content": f"YAML content (raw):\n```yaml\n{content}\n```",
                    "content_type": "yaml",
                    "extension": file_path.suffix,
                    "metadata": {
                        "size_bytes": file_path.stat().st_size
                    }
                }
        except Exception as e:
            log.warning(f"Error processing YAML {file_path}: {e}")
            
            # Fall back to basic text
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
                
            return {
                "title": file_path.name,
                "content": content,
                "content_type": "text",
                "extension": file_path.suffix,
                "metadata": {
                    "size_bytes": file_path.stat().st_size,
                    "error": str(e)
                }
            }
    
    def _handle_docx(self, file_path: Path) -> Optional[Dict[str, Any]]:
        """Handle DOCX file extraction."""
        try:
            if not self.docx_available:
                return None
                
            doc = Document(file_path)
            content = []
            structure = []
            
            # Extract content and structure
            for para in doc.paragraphs:
                content.append(para.text)
                if para.style and "Heading" in para.style.name:
                    level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                    structure.append({
                        "level": level,
                        "title": para.text
                    })
            
            return {
                "title": file_path.name,
                "content": "\n".join(content),
                "content_type": "docx",
                "metadata": {
                    "paragraph_count": len(doc.paragraphs),
                    "section_count": len(doc.sections),
                    "structure": structure
                }
            }
        except Exception as e:
            log.warning(f"Error processing DOCX {file_path}: {e}")
            return {
                "title": file_path.name,
                "content": f"Error processing DOCX file: {e}",
                "content_type": "text",
                "metadata": {"error": str(e)}
            }
            
    def _analyze_json(self, data: Any) -> Dict[str, Any]:
        """
        Analyze JSON structure.
        
        Args:
            data: Parsed JSON data
            
        Returns:
            Dictionary with analysis information
        """
        if isinstance(data, dict):
            # Handle dictionary
            keys = list(data.keys())
            key_types = {k: type(v).__name__ for k, v in data.items()}
            has_nested = any(isinstance(v, (dict, list)) for v in data.values())
            
            return {
                "type": "object",
                "keys_count": len(keys),
                "keys": keys[:20],  # Limit to first 20 keys
                "key_types": key_types,
                "has_nested": has_nested
            }
            
        elif isinstance(data, list):
            # Handle list
            item_count = len(data)
            sample_types = set()
            has_nested = False
            
            # Analyze sample items
            for item in data[:10]:  # Check first 10 items
                sample_types.add(type(item).__name__)
                if isinstance(item, (dict, list)):
                    has_nested = True
                    
            return {
                "type": "array",
                "item_count": item_count,
                "sample_types": list(sample_types),
                "has_nested": has_nested
            }
            
        else:
            # Handle primitive
            return {
                "type": type(data).__name__,
                "value": str(data)
            }
    
    def _create_json_summary(self, data: Any, structure: Dict[str, Any]) -> str:
        """
        Create a summary of JSON data based on structure analysis.
        
        Args:
            data: JSON data
            structure: Structure analysis from _analyze_json
            
        Returns:
            Formatted summary string
        """
        summary = [f"# JSON Data Analysis"]
        
        if structure.get("type") == "object":
            keys = structure.get("keys", [])
            key_count = structure.get("keys_count", 0)
            
            summary.append(f"\n## Structure")
            summary.append(f"- Type: JSON Object")
            summary.append(f"- Keys: {key_count}")
            
            if keys:
                summary.append(f"\n## Keys")
                for i, key in enumerate(keys[:15]):  # Show first 15 keys max
                    key_type = structure.get("key_types", {}).get(key, "unknown")
                    summary.append(f"- `{key}`: {key_type}")
                    
                if len(keys) > 15:
                    summary.append(f"- ... ({len(keys) - 15} more keys)")
                    
            if structure.get("has_nested"):
                summary.append(f"\n## Notes")
                summary.append(f"- Contains nested objects or arrays")
                
        elif structure.get("type") == "array":
            item_count = structure.get("item_count", 0)
            sample_types = structure.get("sample_types", [])
            
            summary.append(f"\n## Structure")
            summary.append(f"- Type: JSON Array")
            summary.append(f"- Items: {item_count}")
            
            if sample_types:
                summary.append(f"\n## Content Types")
                for type_name in sample_types:
                    summary.append(f"- {type_name}")
                    
            if structure.get("has_nested"):
                summary.append(f"\n## Notes")
                summary.append(f"- Contains nested objects or arrays")
                
            # Add sample items if array is not too large
            if item_count > 0 and item_count <= 5:
                summary.append(f"\n## Sample Items")
                for i, item in enumerate(data):
                    summary.append(f"- Item {i+1}: {str(item)[:100]}")
        else:
            # Simple primitive value
            summary.append(f"\n## Content")
            summary.append(f"- Type: {structure.get('type', 'unknown')}")
            summary.append(f"- Value: {structure.get('value', '')}")
            
        return "\n".join(summary)
